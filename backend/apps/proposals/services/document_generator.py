import os
import logging
from datetime import datetime
from pathlib import Path
from django.conf import settings
from django.core.files.base import ContentFile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def ensure_proposals_directory():
    """Ensure the proposals directory exists with correct permissions"""
    proposals_dir = Path(settings.MEDIA_ROOT) / 'proposals'
    
    # Create directory if it doesn't exist
    proposals_dir.mkdir(parents=True, exist_ok=True)
    
    # Set permissions (755 = rwxr-xr-x)
    # Owner: read, write, execute
    # Group/Others: read, execute
    try:
        os.chmod(proposals_dir, 0o755)
    except PermissionError:
        logger.warning(f"Could not set permissions on {proposals_dir}. You may need to run: sudo chown -R $USER:$USER {proposals_dir}")
    
    return proposals_dir

def generate_proposal_document(proposal):
    """
    Generate a DOCX document from a proposal and its sections.
    
    Args:
        proposal: Proposal instance
        
    Returns:
        ContentFile: The generated DOCX file
    """
    # Ensure proposals directory exists with correct permissions
    ensure_proposals_directory()
    
    # Create a new Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading(proposal.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Tender: {proposal.tender.title}")
    doc.add_paragraph(f"Status: {proposal.status.title()}")
    doc.add_paragraph(f"Created: {proposal.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")  # Empty line
    
    # Add sections
    sections = proposal.sections.all().order_by('id')
    
    if not sections.exists():
        doc.add_paragraph("No sections available in this proposal.")
    else:
        for section in sections:
            # Add section heading
            heading = doc.add_heading(section.name, level=1)
            
            # Add section content
            if section.content:
                # Split content by newlines to preserve formatting
                content_lines = section.content.split('\n')
                for line in content_lines:
                    if line.strip():  # Skip empty lines
                        para = doc.add_paragraph(line.strip())
                    else:
                        doc.add_paragraph("")  # Add empty paragraph for spacing
            else:
                doc.add_paragraph("(No content)")
            
            # Add spacing between sections
            doc.add_paragraph("")
    
    # Add page break before appendices if needed
    # (You can extend this to add appendices, tables of contents, etc.)
    
    # Save to in-memory file
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create filename
    filename = f"proposal_{proposal.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # Create ContentFile
    file_content = ContentFile(buffer.read())
    file_content.name = filename
    
    logger.info(f"Generated DOCX document for proposal {proposal.id}")
    
    return file_content

