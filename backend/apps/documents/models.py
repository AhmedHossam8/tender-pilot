from django.db import models
from django.contrib.auth import get_user_model
from apps.projects.models import Project
import pdfplumber
import docx
import os

User = get_user_model()

DOCUMENT_TYPE_CHOICES = [
    ("pdf", "PDF"),
    ("word", "Word"),
    ("excel", "Excel"),
    ("other", "Other"),
]

class ProjectDocument(models.Model):
    """
    Represents a document attached to a project.
    Handles file upload, text extraction, and AI processing metadata.
    """
    project = models.ForeignKey(
        Project,
        on_delete=models.CASCADE,
        related_name="documents"
    )
    file = models.FileField(upload_to="project_attachments/")
    document_type = models.CharField(max_length=20, choices=DOCUMENT_TYPE_CHOICES, default="pdf")

    # AI-related fields
    extracted_text = models.TextField(blank=True, null=True)
    ai_summary = models.TextField(blank=True, null=True)
    ai_processed = models.BooleanField(default=False)
    ai_processed_at = models.DateTimeField(blank=True, null=True)

    # Status and ownership
    is_active = models.BooleanField(default=True)
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name="uploaded_documents",
        null=True,
        blank=True
    )

    # Timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        indexes = [
            models.Index(fields=["ai_processed"]),
            models.Index(fields=["project", "is_active"]),
            models.Index(fields=["created_by"]),
        ]
        constraints = [
            models.UniqueConstraint(fields=["project", "file"], name="unique_document_per_project")
        ]

    def __str__(self):
        return f"{self.project.title} - {self.file.name}"

    def extract_text(self):
        """
        Extract text from the uploaded file based on document type.
        Returns empty string if extraction fails or file is missing.
        """
        if not self.file or not os.path.exists(self.file.path):
            return ""

        try:
            if self.document_type == "pdf":
                return self._extract_pdf_text(self.file.path)
            elif self.document_type in ["word", "docx"]:
                return self._extract_docx_text(self.file.path)
            else:
                return ""
        except Exception as e:
            print(f"Error extracting text from {self.file.path}: {e}")
            return ""

    def _extract_pdf_text(self, file_path):
        """Extract text from a PDF file using pdfplumber."""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()

    def _extract_docx_text(self, file_path):
        """Extract text from a DOCX file using python-docx."""
        text = ""
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text.strip()
